"Operation: Produce MS Word file (docx) from a note and its subnotes."

import io
import json

import docx
from docx.enum.style import WD_STYLE_TYPE
import flask

from operation import BaseOperation

DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

class Operation(BaseOperation):
    "Produce MS Word file (docx) from a note and its subnotes."

    title = "MS Word file"

    def __init__(self, config):
        self.debug = config.get("DEBUG")

    def is_applicable(self, note):
        "Is this operation applicable to the given note?"
        return True

    def get_parameters(self, note):
        "Return the parameters required to control the operation."
        return {
            "subnotes": {
                "type": "checkbox",
                "description": "Include text from subnotes.",
                "label": "Include subnotes.",
                "default": True
            },
            "font_name": {
                "type": "select",
                "description": "The font to use for body text.",
                "values": ["Arial", "New Times Roman", "Calibri"],
                "default": "Arial"
            }
        }

    def execute(self, note, form):
        """Execute the operation for the given note.
        The form is a dictionary containing the required parameters;
        typically 'flask.request.form'.
        Raise ValueError if something is wrong.
        Return True if the note was changed.
        If this operation generates a response, return it.
        Otherwise return None.
        """
        self.document = docx.Document()
        font_name = form.get("font_name")
        if font_name:
            self.document.styles["Normal"].font.name = font_name
            self.document.styles["Title"].font.name = font_name
            for n in range(1, 9):
                self.document.styles[f"Heading {n}"].font.name = font_name
        if form.get("subnotes"):
            notes = list(note.traverse())
        else:
            note.level = 0
            notes = [note]
        for n in notes:
            self.paragraph = self.document.add_paragraph()
            if n.level == 0:
                self.paragraph.style = self.document.styles["Title"]
            else:
                self.paragraph.style = self.document.styles[f"Heading {n.level}"]
            self.run = self.paragraph.add_run(n.title)
            for child in n.ast["children"]:
                self.render(child)
        output = io.BytesIO()
        self.document.save(output)
        response = flask.make_response(output.getvalue())
        response.headers.set("Content-Type", DOCX_MIMETYPE)
        response.headers.set("Content-Disposition", "attachment",
                             filename=f"{note.title}.docx")
        return response

    def render(self, child):
        "Output content of child recursively."
        if child["element"] == "paragraph":
            self.paragraph = self.document.add_paragraph()
            self.paragraph.style = self.document.styles["Normal"]
            self.run = self.paragraph.add_run()
            for child2 in child["children"]:
                self.render(child2)
        elif child["element"] == "raw_text":
            self.run.add_text(child["children"])
        elif child["element"] == "emphasis":
            self.run = self.paragraph.add_run()
            self.run.italic = True
            for child2 in child["children"]:
                self.render(child2)
            self.run = self.paragraph.add_run()
        elif child["element"] == "strong_emphasis":
            self.run = self.paragraph.add_run()
            self.run.bold = True
            for child2 in child["children"]:
                self.render(child2)
            self.run = self.paragraph.add_run()
        elif child["element"] == "heading":
            self.paragraph = self.document.add_paragraph()
            self.paragraph.style = self.document.styles[f"Heading {child['level']}"]
            self.run = self.paragraph.add_run()
            for child2 in child["children"]:
                self.render(child2)
        elif self.debug:
            print("child", json.dumps(child, indent=2))
